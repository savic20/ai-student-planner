"""
File Parsers
------------
Extract text from uploaded PDF and DOCX files.

SUPPORTED FORMATS:
- PDF (.pdf) - Using pypdf
- Word (.docx) - Using python-docx
- Plain text (.txt)

WHY SEPARATE THIS?
- Parsing logic is complex
- Can be tested independently
- Reusable across different agents
"""

import io
import logging
from typing import Optional, Tuple
from pathlib import Path

# PDF parsing
from pypdf import PdfReader

# DOCX parsing
from docx import Document

logger = logging.getLogger(__name__)


class FileParserError(Exception):
    """Custom exception for file parsing errors"""
    pass


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file.
    
    ARGS:
    - file_content: Raw bytes of PDF file
    
    RETURNS:
    Extracted text as string
    
    RAISES:
    FileParserError if parsing fails
    
    EXAMPLE:
    with open("syllabus.pdf", "rb") as f:
        text = extract_text_from_pdf(f.read())
    """
    try:
        # Create PDF reader from bytes
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        # Extract text from all pages
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        if not text_parts:
            raise FileParserError("PDF contains no extractable text")
        
        full_text = "\n\n".join(text_parts)
        
        logger.info(f"Extracted {len(full_text)} characters from PDF ({len(reader.pages)} pages)")
        
        return full_text
    
    except FileParserError:
        raise
    except Exception as e:
        logger.error(f"PDF parsing error: {e}", exc_info=True)
        raise FileParserError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file.
    
    ARGS:
    - file_content: Raw bytes of DOCX file
    
    RETURNS:
    Extracted text as string
    
    EXAMPLE:
    with open("syllabus.docx", "rb") as f:
        text = extract_text_from_docx(f.read())
    """
    try:
        # Create DOCX document from bytes
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        # Extract text from all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        
        if not paragraphs:
            raise FileParserError("DOCX contains no text")
        
        full_text = "\n\n".join(paragraphs)
        
        logger.info(f"Extracted {len(full_text)} characters from DOCX ({len(doc.paragraphs)} paragraphs)")
        
        return full_text
    
    except FileParserError:
        raise
    except Exception as e:
        logger.error(f"DOCX parsing error: {e}", exc_info=True)
        raise FileParserError(f"Failed to parse DOCX: {str(e)}")


def extract_text_from_file(file_content: bytes, filename: str) -> Tuple[str, str]:
    """
    Extract text from uploaded file (auto-detect format).
    
    ARGS:
    - file_content: Raw file bytes
    - filename: Original filename (used to detect type)
    
    RETURNS:
    (extracted_text, file_type)
    
    RAISES:
    FileParserError if format unsupported or parsing fails
    
    EXAMPLE:
    with open("syllabus.pdf", "rb") as f:
        text, file_type = extract_text_from_file(f.read(), "syllabus.pdf")
    print(f"Extracted from {file_type}: {text[:100]}...")
    """
    # Get file extension
    extension = Path(filename).suffix.lower()
    
    # Route to appropriate parser
    if extension == ".pdf":
        text = extract_text_from_pdf(file_content)
        return text, "pdf"
    
    elif extension in [".docx", ".doc"]:
        text = extract_text_from_docx(file_content)
        return text, "docx"
    
    elif extension == ".txt":
        try:
            text = file_content.decode('utf-8')
            logger.info(f"Loaded {len(text)} characters from TXT file")
            return text, "txt"
        except Exception as e:
            raise FileParserError(f"Failed to decode TXT file: {str(e)}")
    
    else:
        raise FileParserError(
            f"Unsupported file format: {extension}. "
            f"Supported formats: .pdf, .docx, .txt"
        )


def validate_syllabus_content(text: str) -> bool:
    """
    Basic validation that extracted text looks like a syllabus.
    
    CHECKS:
    - Minimum length (100 characters)
    - Contains syllabus-like keywords
    
    RETURNS:
    True if content looks valid, False otherwise
    
    USAGE:
    text = extract_text_from_pdf(pdf_bytes)
    if not validate_syllabus_content(text):
        raise ValueError("This doesn't look like a syllabus")
    """
    # Check minimum length
    if len(text.strip()) < 100:
        logger.warning("Text too short to be a syllabus")
        return False
    
    # Check for syllabus-like keywords (case-insensitive)
    text_lower = text.lower()
    syllabus_keywords = [
        "syllabus", "course", "instructor", "assignment",
        "exam", "grade", "schedule", "objectives", "textbook"
    ]
    
    keyword_count = sum(1 for keyword in syllabus_keywords if keyword in text_lower)
    
    if keyword_count < 2:
        logger.warning(f"Only {keyword_count} syllabus keywords found")
        return False
    
    logger.info(f"Syllabus validation passed ({keyword_count} keywords found)")
    return True


# =============================================================================
# USAGE EXAMPLES
# =============================================================================
if __name__ == "__main__":
    # Example 1: Parse PDF
    print("Example 1: Parse PDF")
    try:
        with open("sample_syllabus.pdf", "rb") as f:
            text, file_type = extract_text_from_file(f.read(), "sample_syllabus.pdf")
        print(f"✅ Extracted {len(text)} characters from {file_type}")
        print(f"First 200 chars: {text[:200]}...")
    except FileNotFoundError:
        print("⚠️  sample_syllabus.pdf not found (this is just an example)")
    except FileParserError as e:
        print(f"❌ Parsing failed: {e}")
    
    # Example 2: Validate content
    print("\nExample 2: Validate content")
    test_text = """
    Course Syllabus - Computer Science 101
    Instructor: Dr. Smith
    
    Assignments:
    1. Homework 1 - Due: Sept 15
    2. Midterm Exam - Oct 20
    """
    
    if validate_syllabus_content(test_text):
        print("✅ Content validated")
    else:
        print("❌ Content validation failed")